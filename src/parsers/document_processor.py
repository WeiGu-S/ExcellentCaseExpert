"""
文档处理器实现，支持多种文件格式的处理
"""

import os
import mimetypes
from pathlib import Path
from typing import List, Optional, Tuple, Union
from datetime import datetime

from PIL import Image
import PyPDF2
from docx import Document

from ..models.data_models import ParsedContent, ErrorResponse
from ..utils.constants import (
    SUPPORTED_IMAGE_FORMATS, SUPPORTED_DOCUMENT_FORMATS, SUPPORTED_FORMATS,
    MAX_FILE_SIZE, MAX_IMAGE_SIZE
)
from ..utils.error_handler import ErrorHandler


class DocumentProcessor:
    """文档处理器，支持PNG、JPG、Word、PDF格式"""
    
    def __init__(self):
        """初始化文档处理器"""
        self.error_handler = ErrorHandler()
        self.supported_formats = SUPPORTED_FORMATS
        self.max_file_size = MAX_FILE_SIZE
        self.max_image_size = MAX_IMAGE_SIZE
    
    def validate_file(self, file_path: str) -> Tuple[bool, Optional[str]]:
        """
        验证文件格式和大小
        
        Args:
            file_path: 文件路径
            
        Returns:
            (是否有效, 错误信息)
        """
        try:
            # 检查文件是否存在
            if not os.path.exists(file_path):
                return False, "文件不存在"
            
            # 检查文件是否为文件（不是目录）
            if not os.path.isfile(file_path):
                return False, "路径不是有效的文件"
            
            # 获取文件扩展名
            file_ext = Path(file_path).suffix.lower()
            if not file_ext:
                return False, "文件没有扩展名"
            
            # 检查文件格式是否支持
            if file_ext not in self.supported_formats:
                supported_list = ', '.join(self.supported_formats)
                return False, f"不支持的文件格式。支持的格式: {supported_list}"
            
            # 检查文件大小
            file_size = os.path.getsize(file_path)
            
            # 图片文件大小检查
            if file_ext in SUPPORTED_IMAGE_FORMATS:
                if file_size > self.max_image_size:
                    max_size_mb = self.max_image_size // (1024 * 1024)
                    return False, f"图片文件过大，最大支持 {max_size_mb}MB"
            else:
                # 文档文件大小检查
                if file_size > self.max_file_size:
                    max_size_mb = self.max_file_size // (1024 * 1024)
                    return False, f"文档文件过大，最大支持 {max_size_mb}MB"
            
            # 对于图片文件，额外验证图片格式
            if file_ext in SUPPORTED_IMAGE_FORMATS:
                try:
                    with Image.open(file_path) as img:
                        # 检查图片是否损坏
                        img.verify()
                except Exception as e:
                    return False, f"图片文件损坏或格式无效: {str(e)}"
            
            return True, None
            
        except Exception as e:
            error_msg = f"文件验证失败: {str(e)}"
            return False, error_msg
    
    def get_file_info(self, file_path: str) -> dict:
        """
        获取文件基本信息
        
        Args:
            file_path: 文件路径
            
        Returns:
            文件信息字典
        """
        try:
            file_stat = os.stat(file_path)
            file_path_obj = Path(file_path)
            
            # 获取MIME类型
            mime_type, _ = mimetypes.guess_type(file_path)
            
            info = {
                'name': file_path_obj.name,
                'path': str(file_path_obj.absolute()),
                'size': file_stat.st_size,
                'size_mb': round(file_stat.st_size / (1024 * 1024), 2),
                'extension': file_path_obj.suffix.lower(),
                'mime_type': mime_type,
                'created_time': datetime.fromtimestamp(file_stat.st_ctime),
                'modified_time': datetime.fromtimestamp(file_stat.st_mtime),
            }
            
            # 对于图片文件，获取额外信息
            if info['extension'] in SUPPORTED_IMAGE_FORMATS:
                try:
                    with Image.open(file_path) as img:
                        info.update({
                            'width': img.width,
                            'height': img.height,
                            'format': img.format,
                            'mode': img.mode,
                        })
                except Exception:
                    # 如果无法获取图片信息，继续处理
                    pass
            
            return info
            
        except Exception as e:
            return {'error': f"获取文件信息失败: {str(e)}"}
    
    def process_image(self, file_path: str) -> ParsedContent:
        """
        处理图片文件，准备OCR识别
        
        Args:
            file_path: 图片文件路径
            
        Returns:
            解析后的内容对象
        """
        try:
            # 验证文件
            is_valid, error_msg = self.validate_file(file_path)
            if not is_valid:
                raise ValueError(error_msg)
            
            # 检查是否为图片格式
            file_ext = Path(file_path).suffix.lower()
            if file_ext not in SUPPORTED_IMAGE_FORMATS:
                raise ValueError(f"不是支持的图片格式: {file_ext}")
            
            # 读取并预处理图片
            with Image.open(file_path) as img:
                # 转换为RGB模式（如果需要）
                if img.mode != 'RGB':
                    img = img.convert('RGB')
                
                # 获取图片信息
                width, height = img.size
                
                # 如果图片过大，进行缩放以提高OCR性能
                max_dimension = 2048
                if max(width, height) > max_dimension:
                    ratio = max_dimension / max(width, height)
                    new_width = int(width * ratio)
                    new_height = int(height * ratio)
                    img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
                
                # 保存预处理后的图片到临时位置（如果需要）
                # 这里我们直接返回原始路径，OCR服务会处理
                
                return ParsedContent(
                    text="",  # 图片文件暂时没有文本内容
                    images=[file_path],
                    metadata={
                        'confidence': 1.0,
                        'extraction_time': datetime.now(),
                        'source_file': file_path
                    }
                )
                
        except Exception as e:
            error_response = self.error_handler.handle_error(e, f"处理图片文件: {file_path}")
            raise Exception(error_response.message)
    
    def process_pdf(self, file_path: str) -> ParsedContent:
        """
        处理PDF文件，提取文本内容
        
        Args:
            file_path: PDF文件路径
            
        Returns:
            解析后的内容对象
        """
        try:
            # 验证文件
            is_valid, error_msg = self.validate_file(file_path)
            if not is_valid:
                raise ValueError(error_msg)
            
            # 检查是否为PDF格式
            file_ext = Path(file_path).suffix.lower()
            if file_ext != '.pdf':
                raise ValueError(f"不是PDF格式: {file_ext}")
            
            extracted_text = ""
            page_count = 0
            
            # 读取PDF文件
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                page_count = len(pdf_reader.pages)
                
                # 提取每一页的文本
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            extracted_text += f"\n--- 第{page_num + 1}页 ---\n"
                            extracted_text += page_text + "\n"
                    except Exception as e:
                        # 某一页提取失败，继续处理其他页
                        print(f"PDF第{page_num + 1}页文本提取失败: {e}")
                        continue
            
            # 计算提取置信度（基于提取到的文本量）
            confidence = min(1.0, len(extracted_text.strip()) / 100) if extracted_text.strip() else 0.0
            
            return ParsedContent(
                text=extracted_text.strip(),
                metadata={
                    'confidence': confidence,
                    'extraction_time': datetime.now(),
                    'source_file': file_path
                }
            )
            
        except Exception as e:
            error_response = self.error_handler.handle_error(e, f"处理PDF文件: {file_path}")
            raise Exception(error_response.message)
    
    def process_docx(self, file_path: str) -> ParsedContent:
        """
        处理Word文档，提取文本内容
        
        Args:
            file_path: Word文档路径
            
        Returns:
            解析后的内容对象
        """
        try:
            # 验证文件
            is_valid, error_msg = self.validate_file(file_path)
            if not is_valid:
                raise ValueError(error_msg)
            
            # 检查是否为Word格式
            file_ext = Path(file_path).suffix.lower()
            if file_ext != '.docx':
                raise ValueError(f"不是Word格式: {file_ext}")
            
            # 读取Word文档
            doc = Document(file_path)
            
            extracted_text = ""
            paragraph_count = 0
            
            # 提取段落文本
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    extracted_text += paragraph.text + "\n"
                    paragraph_count += 1
            
            # 提取表格文本
            table_count = 0
            for table in doc.tables:
                extracted_text += f"\n--- 表格{table_count + 1} ---\n"
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        extracted_text += " | ".join(row_text) + "\n"
                table_count += 1
            
            # 计算提取置信度
            confidence = 1.0 if extracted_text.strip() else 0.0
            
            return ParsedContent(
                text=extracted_text.strip(),
                metadata={
                    'confidence': confidence,
                    'extraction_time': datetime.now(),
                    'source_file': file_path
                }
            )
            
        except Exception as e:
            error_response = self.error_handler.handle_error(e, f"处理Word文档: {file_path}")
            raise Exception(error_response.message)
    
    def process_file(self, file_path: str) -> ParsedContent:
        """
        根据文件类型自动选择处理方法
        
        Args:
            file_path: 文件路径
            
        Returns:
            解析后的内容对象
        """
        try:
            # 验证文件
            is_valid, error_msg = self.validate_file(file_path)
            if not is_valid:
                raise ValueError(error_msg)
            
            # 获取文件扩展名
            file_ext = Path(file_path).suffix.lower()
            
            # 根据文件类型选择处理方法
            if file_ext in SUPPORTED_IMAGE_FORMATS:
                return self.process_image(file_path)
            elif file_ext == '.pdf':
                return self.process_pdf(file_path)
            elif file_ext == '.docx':
                return self.process_docx(file_path)
            else:
                raise ValueError(f"不支持的文件格式: {file_ext}")
                
        except Exception as e:
            error_response = self.error_handler.handle_error(e, f"处理文件: {file_path}")
            raise Exception(error_response.message)
    
    def batch_process_files(self, file_paths: List[str]) -> List[Tuple[str, Union[ParsedContent, str]]]:
        """
        批量处理多个文件
        
        Args:
            file_paths: 文件路径列表
            
        Returns:
            处理结果列表，每个元素为(文件路径, 处理结果或错误信息)
        """
        results = []
        
        for file_path in file_paths:
            try:
                parsed_content = self.process_file(file_path)
                results.append((file_path, parsed_content))
            except Exception as e:
                results.append((file_path, str(e)))
        
        return results
    
    def get_supported_formats(self) -> List[str]:
        """获取支持的文件格式列表"""
        return self.supported_formats.copy()
    
    def is_supported_format(self, file_path: str) -> bool:
        """检查文件格式是否支持"""
        file_ext = Path(file_path).suffix.lower()
        return file_ext in self.supported_formats